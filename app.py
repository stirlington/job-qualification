import streamlit as st
import pandas as pd
from datetime import datetime
import os
from docx import Document
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.application import MIMEApplication
import re

# Page configuration
st.set_page_config(page_title="Recruitment Form", layout="wide")

# Display logo
try:
    logo_path = "logo.png"
    st.image(logo_path, width=200)
except Exception as e:
    st.error("Logo not found. Please ensure 'logo.png' is in the same directory as the script.")

# Title
st.title("Job Vacancy Requirements Form")
st.markdown("Please fill out the details for your job vacancy below.")

def is_valid_email(email):
    """Check if email address is valid"""
    pattern = r'^[\w\.-]+@[\w\.-]+\.\w+$'
    return re.match(pattern, email) is not None

def create_word_document(data):
    doc = Document()
    
    # Add heading
    doc.add_heading('Job Vacancy Details', 0)
    
    # Add information
    for key, value in data.items():
        # Skip the submission date and sender email as they're handled separately
        if key not in ["Submission Date", "Sender Email"]:
            p = doc.add_paragraph()
            p.add_run(f"{key}: ").bold = True
            p.add_run(str(value))
    
    # Save the document
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"job_vacancy_{timestamp}.docx"
    doc.save(filename)
    return filename

def send_email(filename, data, sender_email):
    receiver_email = "info@stirlingqr.com"
    
    # Create message
    msg = MIMEMultipart()
    msg['From'] = sender_email
    msg['To'] = receiver_email
    msg['Subject'] = f"New Job Vacancy: {data['Job Title']}"
    
    # Email body
    body = "A new job vacancy has been submitted with the following details:\n\n"
    for key, value in data.items():
        if key != "Sender Email":  # Exclude sender email from the body
            body += f"{key}: {value}\n"
    
    msg.attach(MIMEText(body, 'plain'))
    
    # Attach Word document
    with open(filename, 'rb') as f:
        attachment = MIMEApplication(f.read(), _subtype='docx')
        attachment.add_header('Content-Disposition', 'attachment', filename=filename)
        msg.attach(attachment)
    
    # Send email
    try:
        server = smtplib.SMTP('smtp.gmail.com', 587)
        server.starttls()
        server.login(sender_email, st.session_state.email_password)
        server.send_message(msg)
        server.quit()
        return True
    except Exception as e:
        st.error(f"Error sending email: {str(e)}")
        return False

# Form
with st.form("job_vacancy_form"):
    # Email Information
    st.subheader("Contact Information")
    sender_email = st.text_input("Your Email Address*")
    
    # Basic Job Information
    st.subheader("Basic Job Information")
    job_title = st.text_input("Job Title*")
    department = st.text_input("Department")
    location = st.text_input("Location*")
    
    # Working Arrangements
    st.subheader("Working Arrangements")
    work_model = st.selectbox(
        "Working Model*",
        ["Office-based", "Hybrid", "Remote", "Flexible"]
    )
    if work_model == "Hybrid":
        office_days = st.number_input("Required Office Days per Week", min_value=1, max_value=5)
    
    # Compensation
    st.subheader("Compensation Package")
    col1, col2 = st.columns(2)
    with col1:
        salary_min = st.number_input("Minimum Salary (£)*", min_value=0)
    with col2:
        salary_max = st.number_input("Maximum Salary (£)*", min_value=0)
    
    benefits = st.multiselect(
        "Benefits Package",
        ["Health Insurance", "Dental Insurance", "Life Insurance", "Pension", 
         "Annual Bonus", "Share Options", "Professional Development Budget",
         "Gym Membership", "Private Healthcare", "Mental Health Support"]
    )
    
    bonus_scheme = st.text_area("Bonus Structure Details (if applicable)")
    
    # Qualifications and Requirements
    st.subheader("Qualifications and Requirements")
    experience_years = st.slider("Years of Experience Required", 0, 20, 3)
    education_level = st.selectbox(
        "Minimum Education Level",
        ["None Required", "High School", "Bachelor's Degree", "Master's Degree", "PhD"]
    )
    
    required_skills = st.text_area("Required Skills and Qualifications*")
    preferred_skills = st.text_area("Preferred Skills (Nice to Have)")
    
    # Interview Process
    st.subheader("Interview Process")
    interview_stages = st.multiselect(
        "Interview Stages*",
        ["Phone Screening", "HR Interview", "Technical Interview", 
         "Task/Assignment", "Panel Interview", "Final Interview",
         "Presentation", "Assessment Center"]
    )
    
    interview_details = st.text_area("Additional Interview Process Details")
    
    # Key Success Factors
    st.subheader("Top 3 Most Important Factors")
    success_factor_1 = st.text_input("1st Most Important Factor*")
    success_factor_2 = st.text_input("2nd Most Important Factor*")
    success_factor_3 = st.text_input("3rd Most Important Factor*")
    
    # Additional Information
    st.subheader("Additional Information")
    start_date = st.date_input("Expected Start Date")
    urgent = st.checkbox("This is an urgent requirement")
    additional_notes = st.text_area("Any Additional Information or Special Requirements")
    
    # Submit button
    submitted = st.form_submit_button("Submit Job Vacancy")

# Handle form submission
if submitted:
    # Add sender_email to required fields
    required_fields = [
        (sender_email, "Email Address"),
        (job_title, "Job Title"),
        (location, "Location"),
        (required_skills, "Required Skills"),
        (success_factor_1, "1st Important Factor"),
        (success_factor_2, "2nd Important Factor"),
        (success_factor_3, "3rd Important Factor")
    ]
    
    missing_fields = [field[1] for field in required_fields if not field[0]]
    
    if missing_fields:
        st.error(f"Please fill out the following required fields: {', '.join(missing_fields)}")
    elif not is_valid_email(sender_email):
        st.error("Please enter a valid email address")
    else:
        # Create popup for email password
        email_password = st.text_input("Please enter your email password to send the form", type="password")
        if email_password:
            st.session_state.email_password = email_password
            
            # Create data dictionary
            form_data = {
                "Submission Date": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                "Sender Email": sender_email,
                "Job Title": job_title,
                "Department": department,
                "Location": location,
                "Working Model": work_model,
                "Salary Range": f"£{salary_min:,} - £{salary_max:,}",
                "Benefits": ", ".join(benefits),
                "Bonus Structure": bonus_scheme,
                "Required Experience": f"{experience_years} years",
                "Education Level": education_level,
                "Required Skills": required_skills,
                "Preferred Skills": preferred_skills,
                "Interview Stages": ", ".join(interview_stages),
                "Interview Details": interview_details,
                "Top Factor 1": success_factor_1,
                "Top Factor 2": success_factor_2,
                "Top Factor 3": success_factor_3,
                "Start Date": start_date.strftime("%Y-%m-%d"),
                "Urgent": "Yes" if urgent else "No",
                "Additional Notes": additional_notes
            }
            
            # Create Word document
            doc_filename = create_word_document(form_data)
            
            # Create download button for Word document
            with open(doc_filename, "rb") as file:
                btn = st.download_button(
                    label="Download Job Vacancy Details",
                    data=file,
                    file_name=doc_filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            
            # Send email
            if send_email(doc_filename, form_data, sender_email):
                st.success("Form submitted successfully! Email sent to info@stirlingqr.com")
            else:
                st.warning("Form submitted but there was an issue sending the email. Please download and forward the document manually.")
            
            # Clean up - remove the temporary file
            if os.path.exists(doc_filename):
                os.remove(doc_filename)
            
            # Save to Excel (optional)
            df = pd.DataFrame([form_data])
            file_path = "job_vacancies.xlsx"
            if os.path.exists(file_path):
                existing_df = pd.read_excel(file_path)
                updated_df = pd.concat([existing_df, df], ignore_index=True)
                updated_df.to_excel(file_path, index=False)
            else:
                df.to_excel(file_path, index=False)

# Add some CSS styling
st.markdown("""
    <style>
        .stButton>button {
            width: 100%;
            background-color: #0066cc;
            color: white;
        }
        .stButton>button:hover {
            background-color: #0052a3;
        }
        img {
            display: block;
            margin-left: auto;
            margin-right: auto;
        }
    </style>
    """, unsafe_allow_html=True)
